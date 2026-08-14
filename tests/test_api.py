from __future__ import annotations

import base64
import io
import json
import stat
from urllib.error import HTTPError
from urllib.request import Request, urlopen

from docx import Document


def request_json(base: str, path: str, *, body=None, origin=True, content_type="application/json", key="test-key"):
    raw = None if body is None else json.dumps(body).encode("utf-8")
    headers = {"Accept": "application/json"}
    if raw is not None:
        headers["Content-Type"] = content_type
        headers["Idempotency-Key"] = key
        if origin:
            headers["Origin"] = base
    request = Request(base + path, data=raw, headers=headers, method="POST" if raw is not None else "GET")
    try:
        with urlopen(request) as response:
            return response.status, dict(response.headers), json.load(response)
    except HTTPError as error:
        return error.code, dict(error.headers), json.load(error)


def test_status_redacts_endpoint_and_sets_security_headers(live_server):
    _app, base = live_server
    status, headers, payload = request_json(base, "/api/status")
    assert status == 200
    assert "base_url" not in payload and "api_key" not in payload
    assert "script-src 'self'" in headers["Content-Security-Policy"]
    assert headers["X-Content-Type-Options"] == "nosniff"


def test_model_settings_start_empty_and_persist_http_provider(live_server, kb_root):
    app, base = live_server
    status, _, initial = request_json(base, "/api/settings/model")
    assert status == 200
    assert initial == {
        "configured": False,
        "provider": None,
        "base_url": "",
        "model": None,
        "has_api_key": False,
        "insecure_http": False,
    }

    payload = {
        "provider": "openai-compatible",
        "base_url": "http://models.example.com/v1",
        "api_key": "test-secret-key",
        "model": "test-model",
    }
    saved_status, _, saved = request_json(base, "/api/settings/model", body=payload, key="save-model-settings")
    assert saved_status == 200
    assert saved["configured"] is True
    assert saved["has_api_key"] is True
    assert "api_key" not in saved
    assert app.service.llm.model == "test-model"
    settings_file = kb_root / ".wiki-state" / "model-settings.json"
    assert stat.S_IMODE(settings_file.stat().st_mode) == 0o600
    assert json.loads(settings_file.read_text(encoding="utf-8"))["api_key"] == "test-secret-key"

    status, _, public_status = request_json(base, "/api/status")
    assert status == 200
    assert "base_url" not in public_status and "api_key" not in public_status


def test_model_list_uses_form_values_without_persisting_secret(live_server, monkeypatch):
    app, base = live_server
    captured = {}

    def list_models(provider, base_url, api_key):
        captured.update(provider=provider, base_url=base_url, api_key=api_key)
        return {"models": ["model-a", "model-b"]}

    monkeypatch.setattr(app, "list_models", list_models)
    status, _, payload = request_json(
        base,
        "/api/settings/models",
        body={"provider": "deepseek", "base_url": "http://models.example.com/v1", "api_key": "ephemeral-key"},
        key="load-models",
    )
    assert status == 200
    assert payload == {"models": ["model-a", "model-b"]}
    assert captured["api_key"] == "ephemeral-key"


def test_post_requires_origin_json_and_exact_fields(live_server, snapshot, kb_root):
    _app, base = live_server
    before = snapshot(kb_root)
    status, _, _ = request_json(base, "/api/generate", body={"keyword": "API term"}, origin=False)
    assert status == 403
    status, _, _ = request_json(base, "/api/generate", body={"keyword": "API term"}, content_type="text/plain")
    assert status == 415
    status, _, _ = request_json(base, "/api/generate", body={"keyword": "API term", "unknown": True})
    assert status == 400
    assert snapshot(kb_root) == before


def test_generate_idempotency_replays_one_result(live_server):
    _app, base = live_server
    payload = {"keyword": "API Concept", "from_path": "concepts/base.md", "heading": "", "passage": ""}
    first_status, _, first = request_json(base, "/api/generate", body=payload, key="same-request")
    second_status, _, second = request_json(base, "/api/generate", body=payload, key="same-request")
    assert first_status == 202 and second_status == 200
    assert first["operation_id"] == second["operation_id"]


def test_invalid_meta_does_not_write(live_server, snapshot, kb_root):
    _app, base = live_server
    before = snapshot(kb_root)
    status, _, _ = request_json(base, "/api/meta", body={"path": "concepts/base.md", "category": "nope", "status": "published"})
    assert status == 422
    assert snapshot(kb_root) == before


def test_article_save_accepts_long_markdown_beyond_general_api_limit(live_server, kb_root):
    _app, base = live_server
    status, _, article = request_json(base, "/api/article?path=concepts/base.md")
    assert status == 200
    long_markdown = article["markdown"] + "\n" + ("长篇正文内容。\n" * 20_000)
    assert len(json.dumps({"markdown": long_markdown}).encode("utf-8")) > 64 * 1024

    saved_status, _, saved = request_json(base, "/api/article/save", body={
        "path": article["path"], "markdown": long_markdown,
        "revision": article["revision"], "force": False,
    }, key="save-long-article")

    assert saved_status == 200, saved
    assert saved["conflict"] is False
    assert (kb_root / "wiki" / article["path"]).read_text(encoding="utf-8") == long_markdown


def test_raw_upload_is_json_only_and_idempotent(live_server):
    _app, base = live_server
    payload = {"filename": "material.md", "content": "# Material\n\nRaw content."}
    first_status, _, first = request_json(base, "/api/ingest/upload", body=payload, key="upload-one")
    replay_status, _, replay = request_json(base, "/api/ingest/upload", body=payload, key="upload-one")
    assert first_status == 201 and replay_status == 200
    assert first["raw"]["path"] == "raw/local/material.md"
    assert replay["raw"]["path"] == first["raw"]["path"]
    status, _, _ = request_json(base, "/api/ingest/upload", body={"filename": "bad.png", "content": "x"})
    assert status == 422


def test_binary_document_upload_and_preview(live_server):
    _app, base = live_server
    document = Document()
    document.add_heading("API Word", level=1)
    document.add_paragraph("通过二进制 API 上传的 Word 正文。")
    output = io.BytesIO()
    document.save(output)
    encoded = base64.b64encode(output.getvalue()).decode("ascii")

    status, _, uploaded = request_json(base, "/api/ingest/upload", body={
        "filename": "api-word.docx", "content_base64": encoded,
    }, key="upload-docx")
    assert status == 201
    assert uploaded["raw"]["source_format"] == "DOCX"

    raw_status, _, raw = request_json(base, "/api/raw?path=raw/local/api-word.docx")
    assert raw_status == 200
    assert "通过二进制 API 上传的 Word 正文" in raw["markdown"]


def test_binary_upload_rejects_invalid_base64_without_writing(live_server, snapshot, kb_root):
    _app, base = live_server
    before = snapshot(kb_root)
    status, _, _ = request_json(base, "/api/ingest/upload", body={
        "filename": "bad.pdf", "content_base64": "not-base64!",
    }, key="invalid-base64")
    assert status == 422
    assert snapshot(kb_root) == before
